from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
OUT_PATH = OUT_DIR / "2026-05_研究内容1阶段工作展示.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(92, 105, 117)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_summary(doc)
    add_route(doc)
    doc.add_page_break()
    add_exploration_attempts(doc)
    doc.add_page_break()
    add_oracle_section(doc)
    add_results(doc)
    add_minimum_finding(doc)
    doc.add_page_break()
    add_generalization_section(doc)
    doc.add_page_break()
    add_faq(doc)
    doc.add_page_break()
    add_deliverables(doc)
    add_next_steps(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("2026年5月阶段工作展示  |  第 ")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    add_page_number(paragraph)
    run = paragraph.add_run(" 页")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("基于量子搜索的新能源电力系统机组组合全局优化算法")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("2026年5月研究内容1阶段工作展示")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_BLUE

    add_key_value_table(
        doc,
        [
            ("研究对象", "固定负荷曲线 d 下的 UC/SCUC 启停承诺值函数 oracle"),
            ("搜索变量", "机组启停承诺 x；负荷曲线 d 作为条件量，不进入 Grover 搜索寄存器"),
            ("核心路线", "结构化 max-affine 值函数替代模型 + 定点值寄存器 + 可逆比较器"),
            ("当前验证", "case14，6 台机组，T=3，18 个承诺比特，262144 个候选状态"),
            ("最新结论", "20-bit value-register oracle 支撑 Grover minimum finding 找到与穷举一致的全局最优承诺"),
        ],
        col_widths=(2.0, 4.25),
    )

    add_callout(
        doc,
        "一句话结论",
        "本月完成了从“oracle 逻辑验证”到“Grover minimum finding 小规模闭环仿真”的推进：在固定负荷曲线 d 下，当前 max-affine value-register oracle 能在 case14 T=3 的启停承诺空间中搜索得到与穷举一致的全局最优方案。",
    )


def add_summary(doc: Document) -> None:
    doc.add_heading("1. 本月完成内容概览", level=1)
    add_bullets(
        doc,
        [
            "明确放弃测量读出型 VQC 作为 Grover oracle，转向 value-register + reversible comparator 路线。",
            "完成结构化特征 f(x) 与 max-affine 值函数替代模型：V_hat_theta(x) = max_r (b_r + theta_r^T f(x))。",
            "完成 T=2 和 T=3 的值函数 oracle 标记验证，并针对 T=3 top-128 边界提出 boundary-aware 训练。",
            "完成 20-bit 定点值寄存器量化验证，解释了 16-bit 在边界处失败、20-bit 成功的原因。",
            "新增 Grover minimum finding 状态空间仿真实验，验证 surrogate oracle 能找到真实全局最优启停承诺。",
        ],
    )

    add_small_table(
        doc,
        ["阶段", "主要问题", "本月结果"],
        [
            ["早期 phase/ancilla VQC", "测量或辅助比特泄漏不适合稳定嵌入 Grover", "作为对照和负例保留"],
            ["value-register oracle", "如何相干地写入值函数并阈值比较", "完成 compute、compare、phase mark、uncompute 逻辑"],
            ["max-affine surrogate", "单一线性模型难以表达 UC/SCUC 值函数", "用分段仿射最大值贴近 Benders/割平面结构"],
            ["Grover minimum finding", "Grover 只能找阈值满足态，不能一次直接求最小", "实现阈值迭代仿真，32/32 trials 找到 T=3 全局最优"],
        ],
        [1.35, 2.45, 2.65],
    )


def add_route(doc: Document) -> None:
    doc.add_heading("2. 技术路线演进", level=1)
    add_timeline(
        doc,
        [
            ("5月19日", "建立固定负荷 UC 值函数原型，完成真实 ED/SCUC 成本计算与初代 phase/ancilla oracle 对照实验。"),
            ("5月20日", "转向 value-register/comparator 路线，加入结构化物理特征、定点量化和校准阈值。"),
            ("5月21日", "实现 max-affine 值函数替代模型，并完成 T=2/T=3 标记验证；进一步加入 boundary-aware 训练解决 T=3 top-128。"),
            ("5月22日", "整理研究内容1总结和 README，形成可向导师解释的主线：固定 d、搜索 x、可逆值函数 oracle。"),
            ("5月23日", "完成 Grover minimum finding 闭环仿真，验证 20-bit quantized max-affine oracle 可找到穷举全局最优。"),
        ],
    )


def add_exploration_attempts(doc: Document) -> None:
    doc.add_heading("专题：本月探索尝试与负例沉淀", level=1)
    add_callout(
        doc,
        "为什么要记录负例",
        "本月并不是一开始就得到 max-affine value-register oracle，而是先系统排除了若干看似可行但不适合 Grover 嵌入或泛化不稳的路线。这些负例构成了研究内容1可行性论证的一部分：它们说明当前主线不是偶然选择，而是在可逆性、相干性、排序能力和资源结构之间筛选出来的。",
    )
    add_small_table(
        doc,
        ["尝试方向", "做法", "结果与沉淀"],
        [
            [
                "测量读出型 VQC",
                "先验证 VQC 对值函数的学习能力。",
                "可作为 learnability baseline，但测量会破坏相干性，不能直接作为 Grover oracle。",
            ],
            [
                "threshold phase VQC",
                "训练 U_theta(x,tau) 表示阈值相位。",
                "训练阈值可拟合，但未见阈值泛化不稳，说明直接学相位边界不是主线。",
            ],
            [
                "稀疏 monomial / signed pattern",
                "用高阶 Boolean 项或带正负文字的模板精确标记小规模目标集。",
                "小算例可精确，但接近模式库或查表上界，不能作为可扩展主线。",
            ],
            [
                "物理聚合 / Hamming / bundle 特征",
                "尝试容量、备用、成本 margin、到最优解距离、机组组合模板等特征。",
                "简单聚合特征会漏掉最优或难以覆盖 broader top-k，说明需要更结构化的值函数表示。",
            ],
            [
                "ancilla reversible oracle",
                "构造 O_theta = U_theta^dagger Z_a U_theta，并做状态向量 Grover 仿真。",
                "形式上酉且可逆，但角度不接近 0 或 pi 时会有 ancilla leakage，Grover 放大受影响。",
            ],
            [
                "controlled / explicit two-ancilla",
                "把可行性辅助位和值辅助位显式计算、相位标记、反算。",
                "验证了 feasibility bit 可完全 uncompute，但 value-ancilla leakage 仍是瓶颈。",
            ],
            [
                "leakage reweighting / joint score",
                "对泄漏状态加权训练，并综合 target probability、leakage 和 mark error 选模型。",
                "能降低泄漏，但存在 leakage 与放大概率的折中，仍不如 value-register route 稳定。",
            ],
            [
                "value-register comparator",
                "把值函数写入定点寄存器，再用可逆比较器阈值标记。",
                "解决测量和 leakage 问题，问题转化为值函数近似精度、排序边界和量化位宽。",
            ],
            [
                "structured features",
                "加入容量、备用、启停、转移、同/邻时段交互、merit-order 调度代理。",
                "非查表地提升排序能力，T=2/T=3 均可支持近优目标集标记。",
            ],
            [
                "max-affine / boundary-aware",
                "用 max_r(b_r + theta_r^T f(x)) 表示分段仿射值函数，并强化 top-128 边界。",
                "与 LP 对偶/Benders 割结构一致；T=3 20-bit 下 top-1 至 top-128 全部精确标记。",
            ],
        ],
        [1.45, 2.15, 2.85],
    )
    add_bullets(
        doc,
        [
            "本月负例结论一：直接学习相位或角度，不等于得到 Grover 可用 oracle；相干性、self-inverse、leakage 都必须检查。",
            "本月负例结论二：高阶 Boolean 精确插值可以作为上界参考，但容易退化为模式库，不能作为青年基金主线。",
            "本月正向收敛：value-register + max-affine surrogate 把问题转化为可解释的可逆算术、比较器和阈值排序问题。",
        ],
    )


def add_oracle_section(doc: Document) -> None:
    doc.add_heading("3. 当前 oracle 结构", level=1)
    p = doc.add_paragraph()
    p.add_run("核心值函数定义：").bold = True
    p.add_run(" V_d(x) = startup(x) + min_y C(y; x, d)")

    p = doc.add_paragraph()
    p.add_run("当前电路中表示的是值函数替代模型：").bold = True
    p.add_run(" V_hat_theta(x) = max_r (b_r + theta_r^T f(x))")

    add_code_block(
        doc,
        [
            "|x>|0_f>|0_L>|0_m>|0_c>",
            "-> compute f(x)",
            "-> compute L_r(x) = b_r + theta_r^T f(x)",
            "-> compute m(x) = max_r L_r(x)",
            "-> compare m(x) <= tau and write c",
            "-> phase flip on c",
            "-> uncompute compare / max / affine pieces / features",
            "=> O_tau |x> = (-1)^[V_hat_theta(x) <= tau] |x>",
        ],
    )

    add_callout(
        doc,
        "可逆性说明",
        "整体结构是 U_compute -> phase flip -> U_compute^-1。辅助寄存器最终恢复到 0，只在 |x> 上留下相位，因此该 oracle 是酉变换意义下的可逆 oracle。当前工作验证的是状态空间级逻辑和数值精度，尚未完成硬件级门级综合。",
    )


def add_results(doc: Document) -> None:
    doc.add_heading("4. 主要实验结果", level=1)
    add_small_table(
        doc,
        ["模型", "T", "承诺比特", "有限可行状态", "特征/片段", "20-bit 结果"],
        [
            ["max-affine", "2", "12", "768", "207 / 32", "精确标记 top-1/4/8/16/32/48/64"],
            ["max-affine", "3", "18", "16384", "380 / 32", "精确标记 top-1/4/8/16/32/64，top-128 失败"],
            ["boundary-aware max-affine", "3", "18", "16384", "380 / 32", "精确标记 top-1/4/8/16/32/64/128"],
        ],
        [1.65, 0.45, 0.75, 1.05, 1.0, 2.65],
    )

    add_bullets(
        doc,
        [
            "T=2 max-affine：MAE 约 12.96，max error 约 137.96，20-bit 下全部测试目标集精确标记。",
            "T=3 普通 max-affine：MAE 约 21.70，max error 约 139.38，20-bit 下 top-128 边界失败。",
            "T=3 boundary-aware max-affine：针对 top-128 边界加权训练，calibration margin 从约 -31.84 提升到正裕量，20-bit 下 top-128 精确标记。",
            "16-bit 在 T=3 top-128 仍失败，说明失败原因是边界量化裕量不足，而不是 oracle 可逆结构本身失效。",
        ],
    )

    add_bullets(
        doc,
        [
            "T=3 穷举最优 generator-major bitstring：111000111111111000。",
            "T=3 穷举最优 time-major bitstring：101110101110101110。",
            "T=3 最优总成本：27985.8598204；结果缓存：results/value_cache_case14.json_h3.npz。",
        ],
    )


def add_minimum_finding(doc: Document) -> None:
    doc.add_heading("5. Grover minimum finding 闭环仿真", level=1)
    p = doc.add_paragraph()
    p.add_run("为什么需要 minimum finding：").bold = True
    p.add_run(" Grover oracle 本身只找满足阈值的状态，不能一次直接输出最小值。因此采用阈值迭代思想：当前 incumbent 给出阈值，Grover 搜索预测值更低的候选，测量后由经典外层决定是否更新 incumbent。")

    add_code_block(
        doc,
        [
            "1. classical controller holds incumbent x_best",
            "2. load tau from V_hat_theta(x_best) into threshold register",
            "3. Grover searches states with V_hat_theta(x) < tau",
            "4. measure candidate x",
            "5. classical side evaluates true V_d(x) for validation/update in simulation",
            "6. repeat until no improvement or round limit",
        ],
    )

    add_small_table(
        doc,
        ["oracle 版本", "trials", "成功次数", "mean oracle calls", "mean rounds"],
        [
            ["exact true-value oracle", "32", "32/32", "约 639.16", "约 8.97"],
            ["floating max-affine oracle", "32", "32/32", "约 617.66", "约 8.84"],
            ["20-bit quantized max-affine oracle", "32", "32/32", "约 653.94", "约 8.94"],
        ],
        [2.35, 0.65, 0.9, 1.35, 1.05],
    )

    add_callout(
        doc,
        "准确边界",
        "量子 oracle 内部只比较 V_hat_theta(x) 与 tau；真实 V_d(x) 不在量子电路中求解，只用于离线训练、仿真验证和 classical outer-loop 的 incumbent 更新判断。",
    )


def add_generalization_section(doc: Document) -> None:
    doc.add_heading("专题：泛化性能的数学论证", level=1)
    add_callout(
        doc,
        "核心表述",
        "启停状态 x 作为量子计算基态 |x> 时彼此正交，因此裸状态本身没有连续泛化性；但 UC/SCUC 值函数不是任意离散函数，而是由给定 x 和负荷曲线 d 后的连续调度子问题诱导出来的参数化 LP 值函数。泛化性来自物理特征 f(x,d) 对连续可行域和成本结构的刻画，而不是来自 |x> 与 |x'> 的内积相似。",
    )

    doc.add_heading("A. 裸启停状态不可泛化", level=2)
    add_bullets(
        doc,
        [
            "启停状态空间为 X = {0,1}^n；量子计算基态满足：若 x != x'，则 <x|x'> = 0。",
            "因此，从 Hilbert 空间内积看，两个不同启停状态没有天然接近关系。",
            "No-free-lunch 论证：若训练集 S 不包含 x0，可以构造 V1 和 V2，使它们在 S 上完全相同，但在 x0 上取值差任意大；任何只看 S 的模型都无法同时保证对 V1 和 V2 正确。",
            "结论：如果把启停状态当成无结构离散标签，不能证明泛化。",
        ],
    )

    doc.add_heading("B. UC/SCUC 值函数具有参数化 LP 结构", level=2)
    add_code_block(
        doc,
        [
            "Q(x,d) = min_y c^T y",
            "s.t. A y <= b0 + Bx x + Bd d",
            "",
            "By LP duality:",
            "Q(x,d) = max_lambda lambda^T (b0 + Bx x + Bd d)",
            "       = max_k [a_k + alpha_k^T x + beta_k^T d]",
            "",
            "V(x,d) = startup(x) + Q(x,d)",
        ],
    )
    add_bullets(
        doc,
        [
            "x 通过机组上下限、备用能力、爬坡约束等进入调度可行域；d 通过负荷平衡和备用需求进入右端项。",
            "由强对偶，调度值函数可写成多个仿射函数的逐点最大，因此具有分段线性凸结构。",
            "启停成本 startup(x) 可由启停转移特征表示，例如 startup_g,t = x_g,t * (1 - x_g,t-1)。",
            "因此 max-affine surrogate：V_hat_theta(x,d) = max_r (b_r + theta_r^T f(x,d))，不是黑箱猜测，而是贴近 LP 对偶/Benders 割结构。",
        ],
    )

    doc.add_heading("C. 启停泛化来自特征空间，负荷泛化来自连续参数", level=2)
    add_small_table(
        doc,
        ["对象", "可泛化性的来源", "当前状态"],
        [
            [
                "启停状态 x",
                "不是来自 |x> 的相似性，而是来自容量、备用、启停、转移、交互、merit-order 代理等 f(x,d) 特征对调度结构的刻画。",
                "当前已验证固定 d 下的全状态空间排序/阈值泛化。",
            ],
            [
                "负荷曲线 d",
                "固定 x 时，Q_x(d) = max_k [a_k(x) + beta_k(x)^T d]，关于 d 是分段线性凸函数，并具有 Lipschitz 连续性。",
                "当前实验仍固定 d；跨负荷泛化需把 d 或负荷派生特征加入 f(x,d) 并做多负荷训练/测试。",
            ],
        ],
        [1.25, 3.05, 2.15],
    )

    doc.add_heading("D. oracle 错标风险的可证明条件", level=2)
    add_bullets(
        doc,
        [
            "若真实 top-k 边界间隔为 Delta = V_(k+1) - V_(k)，且 max_x |V_hat(x) - V(x)| < Delta / 2，则 top-k 集合不会被错分。",
            "使用 calibrated threshold 时，只需验证 min_non_target V_hat(x) - max_target V_hat(x) > 0；这个量就是实验中的 calibration margin。",
            "T=3 top-128 的 boundary-aware 训练把 margin 从负值推到正值，正是在证明该阈值 oracle 的排序泛化足以支撑 Grover 标记。",
        ],
    )


def add_faq(doc: Document) -> None:
    doc.add_heading("6. 导师追问准备", level=1)
    rows = [
        (
            "量子电路是否可逆？",
            "是。结构为 U_compute -> phase flip -> U_compute^-1，辅助寄存器最后恢复到 0，因此整体是酉变换。",
        ),
        (
            "20-bit 是不是量子比特不够？",
            "20-bit 只是值寄存器精度，单看 20 个 qubits 并不大。真正资源压力来自完整 oracle 的特征、仿射片段、max 比较和工作位。当前结论是量化精度足够，不是硬件资源已足够。",
        ),
        (
            "需要的辅助比特太多怎么办？",
            "下一步做资源压缩：片段 L_r 串行计算并复用一个片段寄存器；特征逐项生成、累加、反算；用门深换 qubit 数。",
        ),
        (
            "启停状态是否有泛化性？",
            "模型不是查表，而是通过容量、备用、启停、交互和 merit-order 代理等结构化特征学习值函数排序。当前主要证明固定实例全状态空间内的排序/阈值泛化，跨负荷和跨系统泛化是下一步。",
        ),
        (
            "是不是黑箱 AI？",
            "不是黑箱。模型是结构化 max-affine：每个 L_r 是可解释的仿射片段，整体 max_r L_r 对应分段线性值函数结构；oracle 步骤也对应可逆算术和比较器。",
        ),
    ]
    add_small_table(doc, ["问题", "回答要点"], rows, [1.85, 4.6])


def add_deliverables(doc: Document) -> None:
    doc.add_heading("7. 本月代码与结果产出", level=1)
    add_small_table(
        doc,
        ["类型", "文件", "作用"],
        [
            ["核心模块", "qubit_value_function/max_affine.py", "max-affine 值函数替代模型、拟合和诊断"],
            ["核心模块", "qubit_value_function/structured_features.py", "结构化物理特征 f(x)"],
            ["核心模块", "qubit_value_function/grover_minimum.py", "Grover threshold search 与 minimum finding 仿真"],
            ["实验脚本", "experiments/stage1_case14_t2_max_affine_value_surrogate.py", "T=2/T=3 max-affine value-register oracle 实验"],
            ["实验脚本", "experiments/stage1_case14_t3_grover_minimum_finding.py", "T=3 Grover minimum finding 闭环实验"],
            ["结果文件", "results/stage1_case14_t3_grover_minimum_finding.json", "最新 minimum finding 结果"],
            ["文档", "RESEARCH_CONTENT_1_SUMMARY.md / README.md / WORK_LOG.md", "导师沟通、项目入口和工作记录"],
        ],
        [1.0, 2.85, 2.6],
    )


def add_next_steps(doc: Document) -> None:
    doc.add_heading("8. 下一步计划", level=1)
    add_bullets(
        doc,
        [
            "门级资源估计：给出 x 寄存器、值寄存器、特征寄存器、max 比较器和工作位的 qubit count、depth、T-count 粗估。",
            "压缩版 oracle：实现串行片段计算和逐项特征累加，比较并行版本与空间复用版本的资源差异。",
            "降低量化位宽：测试 18-bit、19-bit，并继续增大边界 margin，争取更低位宽下保持 top-128 精确标记。",
            "泛化实验：做 holdout 状态、不同负荷曲线 d、T=4 或不同算例的验证。",
            "申报书落地：把当前路线写成研究内容1的技术方案、阶段目标、可行性证明和风险控制。",
        ],
    )


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [6.25])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    set_cell_margins(cell, 120, 120, 160, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(10.5)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run(body)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [6.25])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    set_cell_margins(cell, 120, 120, 160, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 45, 60)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], col_widths: tuple[float, float]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, list(col_widths))
    for r_idx, (key, value) in enumerate(rows):
        key_cell = table.cell(r_idx, 0)
        val_cell = table.cell(r_idx, 1)
        shade_cell(key_cell, LIGHT_GRAY)
        set_cell_text(key_cell, key, bold=True, color=DARK_BLUE)
        set_cell_text(val_cell, value)
        for cell in [key_cell, val_cell]:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 80, 80, 120, 120)


def add_small_table(doc: Document, headers: list[str], rows: list[list[str] | tuple[str, ...]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, LIGHT_GRAY)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, center=True)
        set_cell_margins(cell, 80, 80, 100, 100)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), center=False)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[idx], 80, 80, 100, 100)


def add_timeline(doc: Document, rows: list[tuple[str, str]]) -> None:
    add_small_table(doc, ["日期", "关键进展"], rows, [1.0, 5.45])


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.2)
    run.font.bold = bold
    run.font.color.rgb = color if color is not None else INK


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int, bottom: int, start: int, end: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        tag = "w:" + edge
        node = margins.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_in: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_in) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)


if __name__ == "__main__":
    main()
